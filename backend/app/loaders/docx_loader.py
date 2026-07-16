"""DOCX 加载器：python-docx，保留段落、表格和图片。"""

from __future__ import annotations

import io
from pathlib import Path

from app.loaders.base import BaseLoader, ImageInfo, LoadedDocument, LoadedPage


class DocxLoader(BaseLoader):
    extensions = (".docx",)

    def load(self, path: str | Path) -> LoadedDocument:
        from docx import Document as DocxDoc
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        p = Path(path)
        doc = DocxDoc(str(p))

        pages: list[LoadedPage] = []
        buf: list[str] = []
        images: list[ImageInfo] = []
        section_idx = 0
        para_count = 0
        img_idx = 0

        # 收集所有图片
        image_map = self._extract_images(doc)

        for para in doc.paragraphs:
            text = (para.text or "").strip()
            style = (para.style.name or "").lower() if para.style else ""

            # 检查段落中的图片
            para_images = self._get_para_images(para, image_map)
            if para_images:
                for img_info in para_images:
                    img_idx += 1
                    img_info.position = len("\n".join(buf))
                    images.append(img_info)
                    buf.append(f"[图片:{img_idx}]")

            if "heading" in style and text:
                if buf:
                    pages.append(
                        LoadedPage(
                            text="\n".join(buf).strip(),
                            page=None,
                            section=section_idx or None,
                            images=images.copy(),
                        )
                    )
                    buf = []
                    images = []
                section_idx += 1
                buf.append(f"## {text}")
            elif text:
                buf.append(text)
                para_count += 1

        for tbl in doc.tables:
            rows = []
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                buf.append("\n".join(rows))

        if buf:
            pages.append(
                LoadedPage(
                    text="\n".join(buf).strip(),
                    page=None,
                    section=section_idx or None,
                    images=images,
                )
            )

        full = "\n\n".join(p.text for p in pages if p.text)
        total_images = sum(len(p.images) for p in pages)

        return LoadedDocument(
            text=full,
            pages=pages,
            metadata={"paragraph_count": para_count, "image_count": total_images},
            source=str(p),
            ext=p.suffix.lower(),
        )

    def _extract_images(self, doc) -> dict:
        """提取文档中所有图片，返回 {rId: (bytes, mime_type)} 映射。"""
        from docx.parts.image import Image as DocxImage

        image_map = {}
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    image_data = rel.target_part.blob
                    # 获取 MIME 类型
                    content_type = rel.target_part.content_type
                    if "png" in content_type:
                        mime = "image/png"
                    elif "jpeg" in content_type or "jpg" in content_type:
                        mime = "image/jpeg"
                    elif "gif" in content_type:
                        mime = "image/gif"
                    elif "bmp" in content_type:
                        mime = "image/bmp"
                    else:
                        mime = content_type or "image/png"
                    image_map[rel.rId] = (image_data, mime)
                except Exception:
                    continue
        return image_map

    def _get_para_images(self, para, image_map: dict) -> list[ImageInfo]:
        """从段落中提取图片信息。"""
        images = []
        try:
            # 遍历段落中的 run，查找图片
            for run in para.runs:
                drawing_elems = run._element.xpath(".//a:blip")
                for blip in drawing_elems:
                    embed = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if embed and embed in image_map:
                        data, mime = image_map[embed]
                        images.append(ImageInfo(image_bytes=data, mime_type=mime))
        except Exception:
            pass
        return images